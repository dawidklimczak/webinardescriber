import streamlit as st
import requests
import json
import time
import os
from openai import OpenAI
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
import tempfile
import PyPDF2

# Konfiguracja API keys z secrets lub zmiennych środowiskowych
def get_api_keys():
    # Próbujemy pobrać z secrets Streamlit Cloud
    try:
        assembly_key = st.secrets["ASSEMBLY_AI_API_KEY"]
        openai_key = st.secrets["OPENAI_API_KEY"]
    except:
        # Jeśli nie ma w secrets, próbujemy ze zmiennych środowiskowych
        assembly_key = os.environ.get("ASSEMBLY_AI_API_KEY")
        openai_key = os.environ.get("OPENAI_API_KEY")
        
        # Jeśli nadal brak, pozwalamy wprowadzić ręcznie
        if not assembly_key or not openai_key:
            if "api_keys_provided" not in st.session_state:
                st.session_state.api_keys_provided = False
            
            if not st.session_state.api_keys_provided:
                with st.form("api_keys_form"):
                    if not assembly_key:
                        assembly_key = st.text_input("Podaj klucz API AssemblyAI", type="password")
                    if not openai_key:
                        openai_key = st.text_input("Podaj klucz API OpenAI", type="password")
                    submit = st.form_submit_button("Zapisz klucze")
                    
                    if submit:
                        if assembly_key and openai_key:
                            st.session_state.api_keys_provided = True
                            st.success("Klucze API zapisane w sesji!")
                        else:
                            st.error("Oba klucze API są wymagane.")
                
                if not st.session_state.api_keys_provided:
                    st.stop()
    
    return assembly_key, openai_key

# Schemat JSON dla analizy webinaru
WEBINAR_ANALYSIS_SCHEMA = {
    "type": "object",
    "properties": {
        "top_quotes": {
            "type": "array",
            "description": "Najciekawsze fragmenty i cytaty do wykorzystania w promocji",
            "items": {"type": "string"},
            "minItems": 5
        },
        "syllabus": {
            "type": "array",
            "description": "Program webinaru/szkolenia w postaci listy tematów z opisami",
            "items": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Tytuł punktu programu"},
                    "description": {"type": "string", "description": "Opis punktu programu na 50 do 100 słów"}
                },
                "required": ["title", "description"]
            },
            "minItems": 3
        },
        "keywords": {
            "type": "array",
            "description": "Słowa kluczowe do promocji",
            "items": {"type": "string"},
            "minItems": 10
        },
        "description": {
            "type": "string",
            "description": "Marketingowy opis szkolenia/webinaru, który może być użyty jako opis produktu (nie zachęta do udziału). Długość: 120-200 słów."
        },
        "benefits": {
            "type": "array",
            "description": "Podpunkty dla sekcji 'Z tego webinaru dowiesz się...'",
            "items": {"type": "string"},
            "minItems": 5,
            "maxItems": 7
        },
        "title": {
            "type": "string",
            "description": "Proponowany chwytliwy tytuł webinaru/szkolenia"
        },
        "target_audience": {
            "type": "string",
            "description": "Dla kogo przeznaczony jest ten webinar/szkolenie"
        },
        "instructor_bio": {
            "type": "string",
            "description": "Proponowany opis prowadzącego bazujący na transkrypcji"
        }
    },
    "required": ["top_quotes", "syllabus", "keywords", "description", "benefits", "title", "target_audience", "instructor_bio"]
}

# Schemat JSON dla analizy ebooka (PDF)
PDF_ANALYSIS_SCHEMA = {
    "type": "object",
    "properties": {
        "top_quotes": {
            "type": "array",
            "description": "Najciekawsze fragmenty i cytaty do wykorzystania w promocji",
            "items": {"type": "string"},
            "minItems": 5
        },
        "main_topics": {
            "type": "array",
            "description": "Główne tematy przewodnie ebooka z opisami",
            "items": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Nazwa tematu przewodniego"},
                    "description": {"type": "string", "description": "Opis tematu na 50 do 100 słów"}
                },
                "required": ["title", "description"]
            },
            "minItems": 3
        },
        "research_references": {
            "type": "array",
            "description": "Fragmenty tekstu odwołujące się do badań naukowych",
            "items": {"type": "string"},
            "minItems": 3
        },
        "keywords": {
            "type": "array",
            "description": "Słowa kluczowe do promocji",
            "items": {"type": "string"},
            "minItems": 10
        },
        "description": {
            "type": "string",
            "description": "Marketingowy opis ebooka, który może być użyty jako opis produktu (nie zachęta do udziału). Długość: 120-200 słów."
        },
        "benefits": {
            "type": "array",
            "description": "Podpunkty dla sekcji 'Z tego ebooka dowiesz się...'",
            "items": {"type": "string"},
            "minItems": 5,
            "maxItems": 7
        },
        "title": {
            "type": "string",
            "description": "Proponowany chwytliwy tytuł ebooka"
        },
        "target_audience": {
            "type": "string",
            "description": "Dla kogo przeznaczony jest ten ebook"
        },
        "author_bio": {
            "type": "string",
            "description": "Proponowany opis autora bazujący na treści ebooka"
        }
    },
    "required": ["top_quotes", "main_topics", "keywords", "description", "benefits", "title", "target_audience", "author_bio"]
}

def transcribe_audio(audio_file, assembly_api_key):
    """Transkrybuje plik audio przy użyciu AssemblyAI."""
    st.info("Rozpoczynam upload pliku do AssemblyAI...")
    
    # Endpoint do wysłania pliku
    upload_endpoint = "https://api.assemblyai.com/v2/upload"
    
    # Nagłówki z kluczem API
    headers = {
        "authorization": assembly_api_key
    }
    
    # Wysyłamy plik
    with open(audio_file, "rb") as f:
        response = requests.post(upload_endpoint, headers=headers, data=f)
    
    if response.status_code != 200:
        st.error(f"Błąd podczas wysyłania pliku: {response.text}")
        return None
    
    audio_url = response.json()["upload_url"]
    st.success("Plik został wysłany. Rozpoczynam transkrypcję...")
    
    # Tworzymy żądanie transkrypcji
    transcript_endpoint = "https://api.assemblyai.com/v2/transcript"
    json_data = {
        "audio_url": audio_url,
        "language_code": "pl"  # Można dostosować do języka webinaru
    }
    
    # Wysyłamy żądanie transkrypcji
    response = requests.post(transcript_endpoint, json=json_data, headers=headers)
    
    if response.status_code != 200:
        st.error(f"Błąd podczas zlecania transkrypcji: {response.text}")
        return None
    
    transcript_id = response.json()["id"]
    st.info(f"Transkrypcja w toku (ID: {transcript_id})...")
    
    # Sprawdzamy status transkrypcji
    polling_endpoint = f"https://api.assemblyai.com/v2/transcript/{transcript_id}"
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    while True:
        response = requests.get(polling_endpoint, headers=headers)
        status = response.json()["status"]
        
        if status == "completed":
            status_text.success("Transkrypcja zakończona!")
            progress_bar.progress(100)
            break
        elif status == "error":
            status_text.error(f"Błąd podczas transkrypcji: {response.json()['error']}")
            return None
        
        status_text.info(f"Status transkrypcji: {status}")
        progress_bar.progress(50)  # Przybliżony progres
        time.sleep(3)
    
    # Pobieramy wynik transkrypcji
    transcript_text = response.json()["text"]
    return transcript_text

def extract_text_from_pdf(pdf_file):
    """Ekstrahuje tekst z pliku PDF."""
    st.info("Ekstrakcja tekstu z pliku PDF...")
    
    try:
        # Otwieramy plik PDF
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        # Inicjalizujemy pusty string na tekst
        text = ""
        
        # Odczyt liczby stron
        num_pages = len(pdf_reader.pages)
        
        # Progres bar
        progress_bar = st.progress(0)
        
        # Iterujemy przez wszystkie strony i wyciągamy tekst
        for i, page in enumerate(pdf_reader.pages):
            text += page.extract_text() + "\n\n"
            # Aktualizacja paska postępu
            progress_bar.progress((i + 1) / num_pages)
        
        st.success("Ekstrakcja tekstu zakończona!")
        return text
    except Exception as e:
        st.error(f"Błąd podczas ekstrakcji tekstu z PDF: {str(e)}")
        return None

def analyze_webinar(text, openai_api_key):
    """Analizuje tekst webinaru przy użyciu OpenAI."""
    st.info("Analizuję tekst webinaru za pomocą OpenAI...")
    
    # Inicjalizacja klienta OpenAI
    openai_client = OpenAI(api_key=openai_api_key)
    
    prompt = f"""
    Przeanalizuj poniższy tekst z webinaru/szkolenia i utwórz szczegółowe materiały marketingowe zgodnie z podanym schematem JSON.
    
    Tekst:
    {text}
    
    Twoja odpowiedź musi zawierać:
    1. "top_quotes": Tablica zawierająca minimum 5 najciekawszych fragmentów i cytatów do wykorzystania w promocji
    2. "syllabus": Tablica obiektów zawierająca program webinaru/szkolenia, gdzie każdy punkt programu ma strukturę:
        {{"title": "Tytuł punktu programu", "description": "Opis punktu programu na 50 do 100 słów"}}
    3. "keywords": Tablica zawierająca minimum 10 słów kluczowych do promocji
    4. "description": String zawierający marketingowy opis szkolenia/webinaru, który może być użyty jako opis produktu (nie zachęta do udziału). Długość: 120-200 słów.
    5. "benefits": Tablica zawierająca 5-7 podpunktów dla sekcji "Z tego webinaru dowiesz się..."
    6. "title": String zawierający proponowany chwytliwy tytuł webinaru/szkolenia
    7. "target_audience": String opisujący dla kogo przeznaczony jest ten webinar/szkolenie
    8. "instructor_bio": String zawierający proponowany opis prowadzącego bazujący na transkrypcji
    
    Zwróć tylko poprawnie sformatowany JSON bez dodatkowego tekstu.
    """
    
    response = openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Jesteś ekspertem od marketingu edukacyjnego, specjalizującym się w analizie i tworzeniu materiałów promocyjnych."},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"},
        temperature=0.7,
    )
    
    try:
        analysis_result = json.loads(response.choices[0].message.content)
        st.success("Analiza zakończona!")
        return analysis_result
    except json.JSONDecodeError:
        st.error("Błąd podczas parsowania odpowiedzi z OpenAI. Odpowiedź nie była poprawnym JSON.")
        st.text(response.choices[0].message.content)
        return None

def analyze_ebook(text, openai_api_key):
    """Analizuje tekst ebooka przy użyciu OpenAI."""
    st.info("Analizuję tekst ebooka za pomocą OpenAI...")
    
    # Inicjalizacja klienta OpenAI
    openai_client = OpenAI(api_key=openai_api_key)
    
    prompt = f"""
    Przeanalizuj poniższy tekst z ebooka i utwórz szczegółowe materiały marketingowe zgodnie z podanym schematem JSON.
    
    Tekst:
    {text}
    
    Twoja odpowiedź musi zawierać:
    1. "top_quotes": Tablica zawierająca minimum 5 najciekawszych fragmentów i cytatów do wykorzystania w promocji
    2. "main_topics": Tablica obiektów zawierająca główne tematy przewodnie ebooka, gdzie każdy temat ma strukturę:
        {{"title": "Nazwa tematu przewodniego", "description": "Opis tematu na 50 do 100 słów"}}
    3. "research_references": Tablica zawierająca minimum 3 fragmenty tekstu odwołujące się do badań naukowych
    4. "keywords": Tablica zawierająca minimum 10 słów kluczowych do promocji
    5. "description": String zawierający marketingowy opis ebooka, który może być użyty jako opis produktu (nie zachęta do udziału). Długość: 120-200 słów.
    6. "benefits": Tablica zawierająca 5-7 podpunktów dla sekcji "Z tego ebooka dowiesz się..."
    7. "title": String zawierający proponowany chwytliwy tytuł ebooka
    8. "target_audience": String opisujący dla kogo przeznaczony jest ten ebook
    9. "author_bio": String zawierający proponowany opis autora bazujący na treści ebooka
    
    Zwróć tylko poprawnie sformatowany JSON bez dodatkowego tekstu.
    """
    
    response = openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Jesteś ekspertem od marketingu edukacyjnego, specjalizującym się w analizie i tworzeniu materiałów promocyjnych."},
            {"role": "user", "content": prompt}
        ],
        response_format={"type": "json_object"},
        temperature=0.7,
    )
    
    try:
        analysis_result = json.loads(response.choices[0].message.content)
        st.success("Analiza zakończona!")
        return analysis_result
    except json.JSONDecodeError:
        st.error("Błąd podczas parsowania odpowiedzi z OpenAI. Odpowiedź nie była poprawnym JSON.")
        st.text(response.choices[0].message.content)
        return None

def create_webinar_document(analysis):
    """Tworzy dokument Word z wynikami analizy webinaru."""
    doc = docx.Document()
    
    # Stylizacja tytułu
    title = doc.add_heading(analysis["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Opis marketingowy
    doc.add_heading("Opis", 1)
    doc.add_paragraph(analysis["description"])
    
    # Dla kogo
    doc.add_heading("Dla kogo", 1)
    doc.add_paragraph(analysis["target_audience"])
    
    # Z tego webinaru dowiesz się...
    doc.add_heading("Z tego webinaru dowiesz się:", 1)
    for benefit in analysis["benefits"]:
        p = doc.add_paragraph()
        p.add_run("• ").bold = True
        p.add_run(benefit)
    
    # Program (syllabus)
    doc.add_heading("Program", 1)
    for i, item in enumerate(analysis["syllabus"], 1):
        # Tytuł punktu programu
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item['title']}").bold = True
        
        # Opis punktu programu
        doc.add_paragraph(item['description']).style = 'List Paragraph'
    
    # Cytaty
    doc.add_heading("Najciekawsze cytaty", 1)
    for quote in analysis["top_quotes"]:
        p = doc.add_paragraph()
        p.add_run(f"❝ {quote} ❞").italic = True
    
    # Słowa kluczowe
    doc.add_heading("Słowa kluczowe", 1)
    p = doc.add_paragraph()
    p.add_run(", ".join(analysis["keywords"]))
    
    # O prowadzącym
    doc.add_heading("O prowadzącym", 1)
    doc.add_paragraph(analysis["instructor_bio"])
    
    # Zapisz dokument w formacie BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def create_ebook_document(analysis):
    """Tworzy dokument Word z wynikami analizy ebooka."""
    doc = docx.Document()
    
    # Stylizacja tytułu
    title = doc.add_heading(analysis["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Opis marketingowy
    doc.add_heading("Opis", 1)
    doc.add_paragraph(analysis["description"])
    
    # Dla kogo
    doc.add_heading("Dla kogo", 1)
    doc.add_paragraph(analysis["target_audience"])
    
    # Z tego ebooka dowiesz się...
    doc.add_heading("Z tego ebooka dowiesz się:", 1)
    for benefit in analysis["benefits"]:
        p = doc.add_paragraph()
        p.add_run("• ").bold = True
        p.add_run(benefit)
    
    # Główne tematy
    doc.add_heading("Główne tematy", 1)
    for i, item in enumerate(analysis["main_topics"], 1):
        # Tytuł tematu
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item['title']}").bold = True
        
        # Opis tematu
        doc.add_paragraph(item['description']).style = 'List Paragraph'
    
    # Odwołania do badań
    doc.add_heading("Odwołania do badań", 1)
    for i, reference in enumerate(analysis["research_references"], 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(reference).italic = True
    
    # Cytaty
    doc.add_heading("Najciekawsze cytaty", 1)
    for quote in analysis["top_quotes"]:
        p = doc.add_paragraph()
        p.add_run(f"❝ {quote} ❞").italic = True
    
    # Słowa kluczowe
    doc.add_heading("Słowa kluczowe", 1)
    p = doc.add_paragraph()
    p.add_run(", ".join(analysis["keywords"]))
    
    # O autorze
    doc.add_heading("O autorze", 1)
    doc.add_paragraph(analysis["author_bio"])
    
    # Zapisz dokument w formacie BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def get_download_link(doc_io, filename="analiza.docx"):
    """Generuje link do pobrania dokumentu Word."""
    b64 = base64.b64encode(doc_io.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Pobierz dokument Word</a>'

def display_webinar_analysis(analysis):
    """Wyświetla wyniki analizy webinaru w interfejsie Streamlit."""
    st.title(analysis["title"])
    
    st.header("Opis marketingowy")
    st.write(analysis["description"])
    
    st.header("Dla kogo")
    st.write(analysis["target_audience"])
    
    st.header("Z tego webinaru dowiesz się...")
    for benefit in analysis["benefits"]:
        st.markdown(f"- {benefit}")
    
    st.header("Program")
    for i, item in enumerate(analysis["syllabus"], 1):
        st.subheader(f"{i}. {item['title']}")
        st.write(item['description'])
    
    st.header("Najciekawsze cytaty")
    for quote in analysis["top_quotes"]:
        st.markdown(f"> *{quote}*")
    
    st.header("Słowa kluczowe")
    st.write(", ".join(analysis["keywords"]))
    
    st.header("O prowadzącym")
    st.write(analysis["instructor_bio"])
    
    # Unikalny identyfikator dla tej sesji wyświetlenia
    display_id = st.session_state.display_count
    
    # Sekcja pobierania
    st.header("Pobierz analizę")
    
    # Tworzenie dokumentu Word
    doc_io = create_webinar_document(analysis)
    st.markdown(get_download_link(doc_io, "analiza_webinaru.docx"), unsafe_allow_html=True)
    
    # Pobierz JSON - dodany unikalny key z identyfikatorem sesji
    st.download_button(
        label="Pobierz dane JSON",
        data=json.dumps(analysis, indent=4, ensure_ascii=False),
        file_name="analiza_webinaru.json",
        mime="application/json",
        key=f"download_json_{display_id}"
    )

def display_ebook_analysis(analysis):
    """Wyświetla wyniki analizy ebooka w interfejsie Streamlit."""
    st.title(analysis["title"])
    
    st.header("Opis marketingowy")
    st.write(analysis["description"])
    
    st.header("Dla kogo")
    st.write(analysis["target_audience"])
    
    st.header("Z tego ebooka dowiesz się...")
    for benefit in analysis["benefits"]:
        st.markdown(f"- {benefit}")
    
    st.header("Główne tematy")
    for i, item in enumerate(analysis["main_topics"], 1):
        st.subheader(f"{i}. {item['title']}")
        st.write(item['description'])
    
    st.header("Odwołania do badań")
    for i, reference in enumerate(analysis["research_references"], 1):
        st.markdown(f"**{i}.** *{reference}*")
    
    st.header("Najciekawsze cytaty")
    for quote in analysis["top_quotes"]:
        st.markdown(f"> *{quote}*")
    
    st.header("Słowa kluczowe")
    st.write(", ".join(analysis["keywords"]))
    
    st.header("O autorze")
    st.write(analysis["author_bio"])
    
    # Unikalny identyfikator dla tej sesji wyświetlenia
    display_id = st.session_state.display_count
    
    # Sekcja pobierania
    st.header("Pobierz analizę")
    
    # Tworzenie dokumentu Word
    doc_io = create_ebook_document(analysis)
    st.markdown(get_download_link(doc_io, "analiza_ebooka.docx"), unsafe_allow_html=True)
    
    # Pobierz JSON - dodany unikalny key z identyfikatorem sesji
    st.download_button(
        label="Pobierz dane JSON",
        data=json.dumps(analysis, indent=4, ensure_ascii=False),
        file_name="analiza_ebooka.json",
        mime="application/json",
        key=f"download_json_{display_id}"
    )

def main():
    st.set_page_config(
        page_title="Analiza Materiałów Edukacyjnych",
        page_icon=None,
        layout="wide"
    )
    
    st.title("Narzędzie do analizy webinarów, szkoleń i e-booków")
    st.write("Wgraj plik audio lub PDF, aby otrzymać analizę i materiały marketingowe")
    
    # Inicjalizacja session_state
    if "display_count" not in st.session_state:
        st.session_state.display_count = 0
    
    # Pobieranie kluczy API
    assembly_api_key, openai_api_key = get_api_keys()
    
    # Wybór typu pliku
    input_type = st.radio(
        "Wybierz typ pliku do analizy:",
        ["Plik audio (webinar/szkolenie)", "Plik PDF (e-book)"]
    )
    
    if input_type == "Plik audio (webinar/szkolenie)":
        uploaded_file = st.file_uploader("Wybierz plik audio (.mp3, .wav, .m4a)", type=["mp3", "wav", "m4a"])
        process_button_label = "Rozpocznij transkrypcję i analizę"
        file_type = "webinar"
    else:
        uploaded_file = st.file_uploader("Wybierz plik PDF", type=["pdf"])
        process_button_label = "Rozpocznij ekstrakcję tekstu i analizę"
        file_type = "ebook"
    
    if uploaded_file is not None:
        # Utworzenie tymczasowego katalogu
        with tempfile.TemporaryDirectory() as temp_dir:
            # Zapisz tymczasowo plik
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            
            # Przyciski akcji
            process_button = st.button(process_button_label)
            
            if process_button:
                # Pozyskaj tekst na podstawie typu pliku
                if file_type == "webinar":
                    text = transcribe_audio(temp_file_path, assembly_api_key)
                    text_source = "Transkrypcja audio"
                else:
                    text = extract_text_from_pdf(uploaded_file)
                    text_source = "Tekst z PDF"
                
                if text:
                    st.session_state.text = text
                    
                    # Pokaż transkrypcję/tekst w expander
                    with st.expander(f"Zobacz pełny {text_source.lower()}"):
                        st.text_area(text_source, text, height=300)
                    
                    # Analizuj tekst w zależności od typu
                    if file_type == "webinar":
                        analysis = analyze_webinar(text, openai_api_key)
                    else:
                        analysis = analyze_ebook(text, openai_api_key)
                    
                    if analysis:
                        st.session_state.analysis = analysis
                        st.session_state.analysis_type = file_type
                        st.session_state.display_count += 1
                        
                        # Wyświetl wyniki analizy
                        if file_type == "webinar":
                            display_webinar_analysis(analysis)
                        else:
                            display_ebook_analysis(analysis)
            
            # Jeśli analiza jest już w session_state, pokaż ją
            elif "analysis" in st.session_state:
                if st.session_state.get("analysis_type") == "webinar":
                    display_webinar_analysis(st.session_state.analysis)
                else:
                    display_ebook_analysis(st.session_state.analysis)

if __name__ == "__main__":
    main()